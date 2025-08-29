"""
Core resume parser module using GLiNER and NuNER Zero models.
Supports dual model extraction with configurable thresholds and entity merging.
"""

import re

import logging
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import Dict, List, Union, Any
from gliner import GLiNER

# Import from our modules
import config
from utils import (
    validate_email,
    normalize_entities,
    merge_adjacent_entities,
    clean_skill_text,
    deduplicate_skills,
    truncate_text
)

# Configure logger
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Advanced resume parser using GLiNER and NuNER Zero models.

    This parser combines both models for optimal entity extraction:
    - GLiNER large-v2.1 for general entity recognition
    - NuNER Zero for superior multi-word entity handling
    """

    def __init__(self):
        """Initialize parser with configured models."""
        self.gliner_model = None
        self.nuner_model = None

        # Load GLiNER model
        if config.MODELS["gliner"]["enabled"]:
            try:
                self.gliner_model = GLiNER.from_pretrained(config.MODELS["gliner"]["model_name"])
                logger.info(f"Loaded GLiNER: {config.MODELS['gliner']['model_name']}")
            except Exception as e:
                logger.error(f"Failed to load GLiNER model: {e}")

        # Load NuNER Zero model
        if config.MODELS["nuner"]["enabled"]:
            try:
                self.nuner_model = GLiNER.from_pretrained(config.MODELS["nuner"]["model_name"])
                logger.info(f"Loaded NuNER Zero: {config.MODELS['nuner']['model_name']}")
            except Exception as e:
                logger.error(f"Failed to load NuNER model: {e}")

        if not self.gliner_model and not self.nuner_model:
            raise RuntimeError("No models loaded successfully")

    def extract_text_from_pdf(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from PDF using PyMuPDF.

        Parameters
        ----------
        file_path : Union[str, Path]
            Path to PDF file

        Returns
        -------
        str
            Extracted text content
        """
        try:
            doc = fitz.open(str(file_path))
            text_parts = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)

            doc.close()

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from PDF: {Path(file_path).name}")
            return full_text

        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from Word document using python-docx.

        Parameters
        ----------
        file_path : Union[str, Path]
            Path to DOCX file

        Returns
        -------
        str
            Extracted text content
        """
        try:
            doc = Document(str(file_path))
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from DOCX: {Path(file_path).name}")
            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from supported file formats.

        Parameters
        ----------
        file_path : Union[str, Path]
            Path to resume file

        Returns
        -------
        str
            Extracted text content
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return ""

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {suffix}")
            return ""

    def extract_entities_with_model(self, text: str, model: GLiNER,
                                    labels: List[str], threshold: float) -> List[Dict]:
        """
        Extract entities using specified model.

        Parameters
        ----------
        text : str
            Input text
        model : GLiNER
            Model instance (GLiNER or NuNER Zero)
        labels : List[str]
            Entity labels to extract
        threshold : float
            Confidence threshold

        Returns
        -------
        List[Dict]
            List of extracted entities
        """
        try:
            # Prepare labels (lowercase for NuNER Zero)
            if model == self.nuner_model and config.MODELS["nuner"].get("requires_lowercase"):
                labels = [label.lower() for label in labels]

            # Extract entities
            entities = model.predict_entities(text, labels, threshold=threshold)

            # Merge adjacent entities for NuNER Zero (handles multi-word entities better)
            if model == self.nuner_model:
                entities = merge_adjacent_entities(entities, text, max_gap=1)

            logger.debug(f"Model extracted {len(entities)} entities with threshold {threshold}")
            return entities

        except Exception as e:
            logger.error(f"Entity extraction failed: {e}")
            return []

    def extract_basic_entities(self, text: str) -> Dict[str, List[str]]:
        """
        Extract basic entities (name, email) using configured approach.

        Parameters
        ----------
        text : str
            Resume text

        Returns
        -------
        Dict[str, List[str]]
            Normalized entities grouped by canonical label
        """
        all_entities = []

        if config.USE_SEPARATE_EXTRACTIONS:
            # Use separate extraction passes for basic entities
            basic_labels = config.ENTITY_LABELS["basic"]

            # GLiNER extraction
            if self.gliner_model:
                gliner_entities = self.extract_entities_with_model(
                    text, self.gliner_model, basic_labels,
                    config.ENTITY_THRESHOLDS["person"]
                )
                all_entities.extend(gliner_entities)

            # NuNER Zero extraction
            if self.nuner_model:
                nuner_entities = self.extract_entities_with_model(
                    text, self.nuner_model, basic_labels,
                    config.ENTITY_THRESHOLDS["person"]
                )
                all_entities.extend(nuner_entities)

        else:
            # Single extraction with all basic labels
            all_basic_labels = config.ENTITY_LABELS["basic"]

            if self.gliner_model:
                entities = self.extract_entities_with_model(
                    text, self.gliner_model, all_basic_labels,
                    config.ENTITY_THRESHOLDS["person"]
                )
                all_entities.extend(entities)

        # Normalize entity labels
        normalized = normalize_entities(all_entities, config.LABEL_NORMALIZATION)
        return normalized

    def extract_skills_entities(self, text: str) -> List[str]:
        """
        Extract skills using configured models and thresholds.

        Parameters
        ----------
        text : str
            Resume text

        Returns
        -------
        List[str]
            Cleaned and deduplicated skills list
        """
        all_entities = []

        if config.USE_SEPARATE_EXTRACTIONS:
            # Separate extraction for skills
            skills_labels = config.ENTITY_LABELS["skills"]

            # GLiNER extraction with skills threshold
            if self.gliner_model:
                gliner_skills = self.extract_entities_with_model(
                    text, self.gliner_model, skills_labels,
                    config.ENTITY_THRESHOLDS["skill"]
                )
                all_entities.extend(gliner_skills)

            # NuNER Zero extraction (better for multi-word skills)
            if self.nuner_model:
                nuner_skills = self.extract_entities_with_model(
                    text, self.nuner_model, skills_labels,
                    config.ENTITY_THRESHOLDS["skill"]
                )
                all_entities.extend(nuner_skills)

        # Fallback: lower threshold if too few skills found
        if len(all_entities) < config.FALLBACK_SETTINGS["min_skills_threshold"]:
            logger.info("Using fallback lower threshold for skills extraction")

            if self.gliner_model:
                fallback_entities = self.extract_entities_with_model(
                    text, self.gliner_model, config.ENTITY_LABELS["skills"],
                    config.FALLBACK_SETTINGS["lower_threshold_fallback"]
                )
                all_entities.extend(fallback_entities)

        # Normalize and clean skills
        normalized = normalize_entities(all_entities, config.LABEL_NORMALIZATION)
        raw_skills = normalized.get("skill", [])

        # Clean individual skills
        cleaned_skills = []
        exclude_words = ["experience", "years", "worked", "using", "with"]

        for skill in raw_skills:
            cleaned_skill = clean_skill_text(skill, exclude_words)
            if cleaned_skill:
                cleaned_skills.append(cleaned_skill)

        # Deduplicate and return
        final_skills = deduplicate_skills(cleaned_skills)
        return final_skills[:25]  # Limit to top 20 skills

    def extract_name_from_entities(self, normalized_entities: Dict[str, List[str]],
                                   text: str) -> str:
        """
        Extract name with fallback heuristics.

        Parameters
        ----------
        normalized_entities : Dict[str, List[str]]
            Normalized entities from models
        text : str
            Original text for fallback extraction

        Returns
        -------
        str
            Extracted name or empty string
        """
        # Try entity extraction results first
        name_candidates = normalized_entities.get("name", [])
        if name_candidates:
            # Filter and return best candidate
            for name in name_candidates:
                name = name.strip()
                words = name.split()
                if 2 <= len(words) <= 4 and len(name) <= 50:
                    return name

        # Fallback: Look in first few lines for name patterns
        lines = text.split('\n')[:3]
        for line in lines:
            line = line.strip()
            words = line.split()

            # Look for 2-4 capitalized words (typical name pattern)
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() for word in words if word.isalpha()):
                    # Avoid common resume headers
                    if not any(header in line.lower() for header in
                               ["resume", "cv", "phone", "email", "@", "address"]):
                        return line

        return ""

    def extract_email_from_entities(self, normalized_entities: Dict[str, List[str]],
                                    text: str) -> str:
        """
        Extract email with regex fallback and validation.

        Parameters
        ----------
        normalized_entities : Dict[str, List[str]]
            Normalized entities from models
        text : str
            Original text for regex fallback

        Returns
        -------
        str
            Validated email or empty string
        """
        # Try entity extraction results first
        email_candidates = normalized_entities.get("email", [])
        for email in email_candidates:
            if validate_email(email, config.EMAIL_VALIDATION):
                return email.lower().strip()

        # Fallback: Regex extraction
        email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        email_matches = email_pattern.findall(text)

        for email in email_matches:
            if validate_email(email, config.EMAIL_VALIDATION):
                return email.lower().strip()

        return ""

    def parse_resume_text(self, text: str) -> Dict[str, Any]:
        """
        Parse resume text and extract structured information.

        Parameters
        ----------
        text : str
            Resume text content

        Returns
        -------
        Dict[str, Any]
            Structured resume data
        """
        if not text.strip():
            return {
                "name": "",
                "email": "",
                "skills": [],
                "metadata": {"error": "Empty text"}
            }

        # Truncate text if too long (for processing efficiency)
        max_length = config.FALLBACK_SETTINGS.get("max_text_length", 5000)
        if len(text) > max_length:
            text = truncate_text(text, max_length)
            logger.info(f"Truncated text to {max_length} characters")

        # Extract basic entities (name, email)
        basic_entities = self.extract_basic_entities(text)

        # Extract skills
        skills = self.extract_skills_entities(text)

        # Process name and email with fallbacks
        name = self.extract_name_from_entities(basic_entities, text)
        email = self.extract_email_from_entities(basic_entities, text)

        # Build result
        result = {
            "name": name,
            "email": email,
            "skills": skills,
            "metadata": {
                "text_length": len(text),
                "models_used": [],
                "entities_found": len(basic_entities.get("name", [])) +
                                  len(basic_entities.get("email", [])) + len(skills),
                "processing_successful": bool(name or email or skills)
            }
        }

        # Track which models were used
        if self.gliner_model:
            result["metadata"]["models_used"].append("GLiNER")
        if self.nuner_model:
            result["metadata"]["models_used"].append("NuNER_Zero")

        logger.info(f"Parsing complete - Name: {'y' if name else '✗'}, "
                    f"Email: {'y' if email else '✗'}, Skills: {len(skills)}")

        return result

    def parse_resume_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse resume file (PDF or DOCX) and extract structured information.

        Parameters
        ----------
        file_path : Union[str, Path]
            Path to resume file

        Returns
        -------
        Dict[str, Any]
            Structured resume data
        """
        file_path = Path(file_path)
        logger.info(f"Processing file: {file_path.name}")

        # Extract text from file
        text = self.extract_text_from_file(file_path)

        if not text:
            return {
                "name": "",
                "email": "",
                "skills": [],
                "metadata": {
                    "error": f"Failed to extract text from {file_path.name}",
                    "file_path": str(file_path)
                }
            }

        # Parse extracted text
        result = self.parse_resume_text(text)
        result["metadata"]["file_path"] = str(file_path)
        result["metadata"]["file_format"] = file_path.suffix.lower()

        return result


# Factory function for easy parser creation
def create_parser() -> ResumeParser:
    """
    Factory function to create and return a configured ResumeParser.

    Returns
    -------
    ResumeParser
        Initialized parser instance
    """
    return ResumeParser()


